#!/usr/bin/env python3
"""生成国企/机关标准格式Word文档"""

from docx import Document
from docx.shared import Pt, Cm, Mm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ===== 页面设置：GB/T 9704-2012 公文标准 =====
section = doc.sections[0]
section.page_width = Cm(21.0)       # A4
section.page_height = Cm(29.7)
section.top_margin = Cm(3.7)        # 上37mm
section.bottom_margin = Cm(3.5)     # 下35mm
section.left_margin = Cm(2.8)       # 左28mm
section.right_margin = Cm(2.6)      # 右26mm


def set_line_spacing(paragraph, spacing_pt):
    """设置固定行距"""
    pPr = paragraph._element.get_or_add_pPr()
    spacing = pPr.find(qn('w:spacing'))
    if spacing is None:
        spacing = OxmlElement('w:spacing')
        pPr.append(spacing)
    spacing.set(qn('w:line'), str(int(spacing_pt * 20)))  # 单位：1/20磅
    spacing.set(qn('w:lineRule'), 'exact')


def add_paragraph(text, font_name, font_size_pt, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER,
                  first_line_indent=None, space_before=0, space_after=0):
    """添加段落，统一设置字体、字号、对齐"""
    p = doc.add_paragraph()
    p.alignment = align

    # 段前段后
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)

    # 首行缩进
    if first_line_indent:
        p.paragraph_format.first_line_indent = Pt(first_line_indent)

    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size_pt)
    run.font.bold = bold
    # 设置中文字体（用于兼容）
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

    # 行距固定
    set_line_spacing(p, 28)

    return p


# ===== 正文内容 =====

# 红色分隔线（标准公文头）
# 标准公文头一般是一行红色分隔线，我们用 # 模拟
add_paragraph("━" * 30, "黑体", 16, bold=False,
              align=WD_ALIGN_PARAGRAPH.CENTER)

# 发文机关
add_paragraph("青云科技集团文件", "黑体", 16, bold=True,
              align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)

# 文号
add_paragraph("青云发〔2026〕88号", "仿宋", 16, bold=False,
              align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

# 红色分隔线
add_paragraph("━" * 30, "黑体", 16, bold=False,
              align=WD_ALIGN_PARAGRAPH.CENTER)

# 标题（二号小标宋 ≈ 22pt）
add_paragraph("关于进一步提升投标工作质量的整改方案", "黑体", 22, bold=True,
              align=WD_ALIGN_PARAGRAPH.CENTER, space_before=12, space_after=12)

# 主送机关
add_paragraph("集团各部门、各区域分公司：", "仿宋", 16, bold=False,
              align=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=32)

# 正文段落
body_text = (
    "根据集团2026年度第3次经营分析会议精神，结合近期振石集团投标项目失标的深刻教训，"
    "为进一步规范投标工作流程，提升投标质量与中标率，现就有关工作提出以下整改要求："
)
add_paragraph(body_text, "仿宋", 16, bold=False,
              align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line_indent=32)

body_text2 = (
    "一、严格落实标前评审制度。所有投标项目必须在投标截止日前至少15个工作日完成标前评审，"
    "评审内容包括技术方案可行性、商务报价合理性、交付能力匹配度及风险评估等。评审通过后方可投递标书。"
)
add_paragraph(body_text2, "仿宋", 16, bold=False,
              align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line_indent=32)

body_text3 = (
    "二、加强演示环节全流程演练。技术演示是决定投标成败的关键环节，"
    "各区域分公司须在正式演示前进行不少于3次内部模拟演练，确保业务流程跑通、"
    "演示环境稳定、应答口径统一。"
)
add_paragraph(body_text3, "仿宋", 16, bold=False,
              align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line_indent=32)

body_text4 = (
    "三、建立投标复盘机制。凡未中标的投标项目，须在收到结果通知后7个工作日内完成复盘分析，"
    "形成书面复盘报告，报集团投标管理委员会备案。复盘结果纳入部门绩效考核。"
)
add_paragraph(body_text4, "仿宋", 16, bold=False,
              align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line_indent=32)

# 落款
today = datetime.date.today()
year_cn = ["〇", "一", "二", "三", "四", "五", "六", "七", "八", "九", "十",
           "十一", "十二"]
month = year_cn[today.month]
day = year_cn[today.day]

# 空行
add_paragraph("", "仿宋", 16, space_before=24)

# 落款单位
add_paragraph("青云科技集团投标管理委员会", "仿宋", 16, bold=False,
              align=WD_ALIGN_PARAGRAPH.RIGHT)
add_paragraph(f"{today.year}年{today.month}月{today.day}日", "仿宋", 16, bold=False,
              align=WD_ALIGN_PARAGRAPH.RIGHT)

# 保存
output_path = "/root/obsidian-vault/星期五的工作区/关于进一步提升投标工作质量的整改方案.docx"
doc.save(output_path)
print(f"✅ 文档已生成: {output_path}")
